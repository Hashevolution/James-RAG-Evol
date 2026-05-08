"""Document export — render a JAMES answer to .md / .docx / .txt.

Item #4 (2026-05-08 user feedback):
  "문서 등 양식에 맞는 파일로 제시하라는 명령어가 작동 가능?"

Scope at v0.2:
  - .md   — always works (text passthrough; native JAMES output is
            already mostly markdown so no transformation needed)
  - .txt  — always works (strips markdown markers for plain prose)
  - .docx — works when python-docx is importable; otherwise falls
            back to .md output and the response_dict carries a
            `fallback_reason` so the UI can show why
  - .pdf  — NOT supported in v0.2 (would need reportlab/weasyprint
            which has heavy native deps); falls back to .md

Why bytes-in-memory rather than a temp file
  The export endpoint sends FileResponse with the bytes directly.
  No temp file = no leak surface, no /tmp filling up under load,
  no race between caller download + cleanup.

Why not import python-docx at module load
  Lazy import inside `_export_docx`. Cold-start performance for the
  normal `.md` / `.txt` path stays free. ImportError handled
  explicitly (falls back to md, not crash).

Filename safety
  `_safe_filename` strips path traversal characters + clamps length.
  Caller-supplied filename is treated as untrusted (operator could
  pass anything via the API).
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass
from datetime import datetime
from typing import Tuple


SUPPORTED_FORMATS = ("md", "txt", "docx", "py")
PDF_DEFERRED_NOTE = (
    "PDF export is deferred to v0.3 — heavy native deps. "
    "Returning markdown instead."
)


@dataclass(frozen=True)
class ExportResult:
    """Returned by `export_document`. Caller wires `data` into the
    HTTP response, `mime` into Content-Type, `filename` into the
    Content-Disposition attachment hint.

    `actual_format` may differ from the requested format when a
    fallback fired (e.g. requested docx, python-docx unavailable,
    fell back to md). `fallback_reason` is a human-readable note
    the UI can surface so the operator knows WHY.
    """
    data:            bytes
    mime:            str
    filename:        str
    actual_format:   str
    fallback_reason: str = ""


# ─── Filename hygiene ────────────────────────────────────────────

_FILENAME_BAD_CHARS = re.compile(r"[\x00-\x1f<>:\"/\\|?*]+")


def _safe_filename(raw: str, default_stem: str = "james_answer") -> str:
    """Strip path traversal + control chars; clamp to 80 chars.

    Returns just the stem (caller appends the extension based on
    actual_format). An all-bad-character input produces the default.
    """
    stem = (raw or "").strip()
    # Drop any path component the caller might have included.
    stem = stem.replace("\\", "/").rsplit("/", 1)[-1]
    # Drop a user-supplied extension — we own the format mapping.
    stem = re.sub(r"\.[A-Za-z0-9]{1,5}$", "", stem)
    stem = _FILENAME_BAD_CHARS.sub("_", stem)
    stem = stem.strip("._ ")
    if not stem:
        stem = default_stem
    return stem[:80]


# ─── Format-specific renderers ───────────────────────────────────

def _export_md(content: str) -> bytes:
    """Pass-through. JAMES native output is already markdown-flavored
    (the natural-flow prompt produces prose, not raw HTML)."""
    return content.encode("utf-8")


def _export_txt(content: str) -> bytes:
    """Strip the most common markdown markers so .txt is clean prose.

    Conservative — leaves URLs, parentheses, etc untouched. The goal
    is "looks reasonable in Notepad", not full markdown→plaintext."""
    text = content
    # Headers: "# Title" / "## Sub" → "Title" / "Sub"
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # Bold/italic markers around words: **x** / *x* / __x__ / _x_
    text = re.sub(r"(\*\*|__)(.+?)\1", r"\2", text)
    text = re.sub(r"(\*|_)(.+?)\1", r"\2", text)
    # Inline code: `x` → x
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Bullets: "- " / "* " → "• "
    text = re.sub(r"^\s*[-*]\s+", "• ", text, flags=re.MULTILINE)
    return text.encode("utf-8")


def _export_py(content: str) -> bytes:
    """Render a chat answer to a .py file.

    User scenario: in coding mode the LLM produces an answer with one
    or more ```python ... ``` fences. They want to save just the code
    (not the prose around it) as runnable .py.

    Strategy:
      - Pull out ```python / ```py / unspecified-language fenced blocks.
      - Concatenate in order, with a single blank line between blocks.
      - Prepend a header comment recording the export timestamp +
        the source so the file is self-describing.
      - If the answer has no fenced blocks at all, write the entire
        prose as `#`-prefixed comments — preserves the answer text
        as a documentation-only `.py`. The file still imports cleanly
        (just does nothing) so it's a safe fallback.

    No syntax validation — running the code is the user's call. We
    don't want to swallow code that contains pseudo-code or partial
    snippets the user explicitly asked for.
    """
    code_re = re.compile(r"```(?:python|py)?\s*\n([\s\S]*?)```", re.IGNORECASE)
    blocks = [m.group(1).strip() for m in code_re.finditer(content)]
    blocks = [b for b in blocks if b]
    header = (
        f"# Exported from JAMES at {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}\n"
        f"# Source: chat answer\n\n"
    )
    if blocks:
        body = "\n\n".join(blocks) + "\n"
    else:
        # No fenced code → escape prose as comments. Lines longer than
        # 200 chars get split so flake8/pylint don't flag the resulting
        # file with line-too-long on every line.
        lines = []
        for raw_line in content.splitlines():
            line = raw_line.rstrip()
            if not line:
                lines.append("#")
                continue
            while len(line) > 197:
                lines.append("# " + line[:197])
                line = line[197:]
            lines.append("# " + line)
        body = "\n".join(lines) + "\n"
    return (header + body).encode("utf-8")


def _export_docx(content: str) -> Tuple[bytes, str]:
    """Render content to a .docx via python-docx. On ImportError
    return (md_bytes, fallback_reason) instead of raising — keeps
    the export path resilient to missing optional deps."""
    try:
        from docx import Document  # type: ignore
    except ImportError as e:
        return (
            _export_md(content),
            f"python-docx not installed ({e}); returning markdown instead. "
            f"Install with: pip install python-docx"
        )

    doc = Document()
    # Render line-by-line. Each line that begins with `# ... ` gets
    # mapped to a Word heading; bullets and quoted lines get their
    # own styles. Plain lines are paragraphs. Conservative — this
    # is a v0.2 first cut, not a full markdown engine.
    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        h = re.match(r"^(#{1,6})\s+(.+)$", line)
        if h:
            level = min(len(h.group(1)), 4)
            doc.add_heading(h.group(2), level=level)
            continue
        bul = re.match(r"^\s*[-*•]\s+(.+)$", line)
        if bul:
            doc.add_paragraph(bul.group(1), style="List Bullet")
            continue
        if line.startswith(">"):
            doc.add_paragraph(line.lstrip("> ").strip(), style="Intense Quote")
            continue
        doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue(), ""


# ─── Public entry point ──────────────────────────────────────────

_MIME_BY_FORMAT = {
    "md":   "text/markdown; charset=utf-8",
    "txt":  "text/plain; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "py":   "text/x-python; charset=utf-8",
}

_EXT_BY_FORMAT = {"md": ".md", "txt": ".txt", "docx": ".docx", "py": ".py"}


def export_document(
    content:  str,
    format:   str = "md",
    filename: str = "",
) -> ExportResult:
    """Render `content` as the requested format.

    Args:
      content:  the answer text to export. Pass empty string and
                you get an empty file (no error — caller's choice).
      format:   one of `md` / `txt` / `docx`. Unknown values fall
                back to `md` with a fallback_reason. `pdf` is
                explicitly noted as v0.3+ work.
      filename: optional stem. Sanitized; default is
                `james_answer_<UTC-stamp>` if empty/invalid.

    Returns:
      ExportResult with bytes, MIME, sanitized filename (with the
      correct extension applied based on `actual_format`), and a
      `fallback_reason` when the format was downgraded.
    """
    requested = (format or "md").strip().lower()
    actual = requested
    fallback_reason = ""

    if requested == "pdf":
        actual = "md"
        fallback_reason = PDF_DEFERRED_NOTE
    elif requested not in SUPPORTED_FORMATS:
        actual = "md"
        fallback_reason = (
            f"unsupported format {requested!r}; supported: "
            f"{', '.join(SUPPORTED_FORMATS)}. Returning markdown."
        )

    # Resolve stem.
    if not filename:
        stem = f"james_answer_{datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')}"
    else:
        stem = _safe_filename(filename)

    if actual == "md":
        data = _export_md(content)
    elif actual == "txt":
        data = _export_txt(content)
    elif actual == "docx":
        data, docx_fallback = _export_docx(content)
        if docx_fallback:
            # python-docx missing → degrade to md, keep both reasons.
            actual = "md"
            fallback_reason = docx_fallback
    elif actual == "py":
        data = _export_py(content)
    else:
        # Defensive — should never reach here given the prelude.
        data = _export_md(content)
        actual = "md"

    return ExportResult(
        data=data,
        mime=_MIME_BY_FORMAT[actual],
        filename=stem + _EXT_BY_FORMAT[actual],
        actual_format=actual,
        fallback_reason=fallback_reason,
    )
