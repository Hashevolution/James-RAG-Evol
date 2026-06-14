"""Render formatted text to downloadable output bytes.

Supported formats (v1.1): ``md`` / ``txt`` / ``html`` / ``docx``.

The HTML wrapper is minimal and fully escaped — no inline JavaScript,
no external resources — so an operator can open the output safely. The
``docx`` path uses ``python-docx`` (Apache-2.0, lazy import) so a
Hangul / Microsoft Word / 한컴오피스 user can open and re-save the
output natively (한글 2010+ opens .docx directly; "다른 이름으로 저장"
→ .hwp is a one-click step). See
``docs/design/v0.6-template-formatting-ui.md`` §6.

Markdown → DOCX conversion is intentionally light:

  * lines starting with ``# `` / ``## `` / ``### `` become
    headings (level 1 / 2 / 3),
  * ``**bold**`` runs become bold text,
  * everything else is a paragraph in the default style.

The goal is "Hangul-importable structured output", not a full Markdown
renderer — operators paste the result into 한글 / Word and edit.
"""
from __future__ import annotations

import html as _html
import io
import re

VALID_FORMATS = ("md", "txt", "html", "docx")

_EXT = {"md": ".md", "txt": ".txt", "html": ".html", "docx": ".docx"}

_HTML_SHELL = (
    "<!DOCTYPE html>\n"
    "<html lang=\"en\">\n<head>\n<meta charset=\"UTF-8\">\n"
    "<meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">\n"
    "<title>{title}</title>\n"
    "<style>body{{font-family:system-ui,sans-serif;max-width:48rem;"
    "margin:2rem auto;padding:0 1rem;line-height:1.6}}"
    "pre{{white-space:pre-wrap;word-wrap:break-word}}</style>\n"
    "</head>\n<body>\n<pre>{body}</pre>\n</body>\n</html>\n"
)

# Bold-run pattern for the docx pass. Non-greedy, single line.
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def extension_for(fmt: str) -> str:
    """Return the file extension (with dot) for ``fmt``."""
    if fmt not in _EXT:
        raise ValueError(f"unsupported format {fmt!r}; valid: {VALID_FORMATS}")
    return _EXT[fmt]


def _heading_level(line: str) -> int:
    """Return 1/2/3 for ``# `` / ``## `` / ``### `` prefixes, else 0."""
    if line.startswith("### "):
        return 3
    if line.startswith("## "):
        return 2
    if line.startswith("# "):
        return 1
    return 0


def _add_runs_with_bold(paragraph, text: str) -> None:
    """Add ``text`` to ``paragraph`` honouring ``**bold**`` runs."""
    last = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def _render_docx(formatted_text: str, title: str) -> bytes:
    """Render Markdown-ish text to a .docx file. ``python-docx`` is lazy-
    imported so the rest of the templating module loads on installs that
    haven't added the dependency yet (clean ImportError message)."""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "docx output requires python-docx; install with "
            "`pip install python-docx`"
        ) from e

    doc = Document()
    if title:
        # Use the document's built-in title style — visible to the
        # 한글/Word reader as the top-level heading.
        doc.add_heading(title, level=0)

    # Split on blank lines so a stretch of contiguous lines becomes one
    # paragraph (mirrors a casual Markdown reader's mental model).
    block: list[str] = []

    def _flush():
        if not block:
            return
        first = block[0]
        lvl = _heading_level(first)
        if lvl > 0 and len(block) == 1:
            doc.add_heading(first[lvl + 1:], level=lvl)
        else:
            p = doc.add_paragraph()
            for i, ln in enumerate(block):
                if i > 0:
                    p.add_run("\n")
                _add_runs_with_bold(p, ln)
        block.clear()

    for raw_line in (formatted_text or "").splitlines():
        line = raw_line.rstrip()
        if not line:
            _flush()
            continue
        # A heading line always flushes the prior block, then is its own
        # paragraph (one heading per add_heading call).
        if _heading_level(line) > 0:
            _flush()
            block.append(line)
            _flush()
            continue
        block.append(line)
    _flush()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render(formatted_text: str, fmt: str = "md", *, title: str = "document") -> bytes:
    """Render ``formatted_text`` to bytes in the requested format.

    - ``md`` / ``txt``: the text verbatim, UTF-8 encoded.
    - ``html``: the text HTML-escaped inside a minimal, JS-free shell.
    - ``docx``: light Markdown → DOCX conversion (headings + bold runs).
    """
    if fmt not in VALID_FORMATS:
        raise ValueError(f"unsupported format {fmt!r}; valid: {VALID_FORMATS}")
    text = formatted_text or ""
    if fmt == "html":
        doc = _HTML_SHELL.format(
            title=_html.escape(title or "document"),
            body=_html.escape(text),
        )
        return doc.encode("utf-8")
    if fmt == "docx":
        return _render_docx(text, title or "document")
    return text.encode("utf-8")


__all__ = ["VALID_FORMATS", "extension_for", "render"]
